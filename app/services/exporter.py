import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── DOCX export ─────────────────────────────────────────────────────────────

def export_docx(optimized_text: str) -> bytes:
    """Convert plain resume text to a formatted, editable Word .docx document."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # Remove default empty paragraph Word adds
    for p in list(doc.paragraphs):
        _remove_para(p)

    lines = optimized_text.split("\n")
    first_content_line = True

    for line in lines:
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph()
            continue

        # First non-empty line → name (large, bold, centered)
        if first_content_line:
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(16)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)
            first_content_line = False
            continue

        # All-caps short line → section heading with underline border
        if stripped.isupper() and len(stripped) < 60:
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(2)
            _add_bottom_border(p)
            continue

        # Bullet point
        if stripped.startswith(("•", "-", "*")):
            bullet_text = stripped.lstrip("•-* ").strip()
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(bullet_text)
            run.font.size = Pt(10.5)
            p.paragraph_format.space_after = Pt(1)
            continue

        # Regular line
        p = doc.add_paragraph()
        run = p.add_run(stripped)
        run.font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(1)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _remove_para(para):
    p = para._element
    p.getparent().remove(p)


def _add_bottom_border(para):
    pPr   = para._p.get_or_add_pPr()
    pBdr  = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1a1a1a")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── PDF export (reportlab — pure Python, no GTK needed) ──────────────────────

def export_pdf(optimized_text: str) -> bytes:
    """Convert plain resume text to PDF using reportlab."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()
    name_style = ParagraphStyle("Name", fontSize=16, fontName="Helvetica-Bold",
                                alignment=TA_CENTER, spaceAfter=6)
    heading_style = ParagraphStyle("Heading", fontSize=11, fontName="Helvetica-Bold",
                                   spaceBefore=10, spaceAfter=3,
                                   borderPadding=(0, 0, 2, 0),
                                   underlineWidth=0.5)
    body_style = ParagraphStyle("Body", fontSize=10.5, fontName="Helvetica",
                                spaceAfter=2, leading=14)
    bullet_style = ParagraphStyle("Bullet", fontSize=10.5, fontName="Helvetica",
                                  leftIndent=16, spaceAfter=2, leading=14,
                                  bulletIndent=6)

    story = []
    lines = optimized_text.split("\n")
    first_content_line = True

    for line in lines:
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 4))
            continue

        safe = stripped.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        if first_content_line:
            story.append(Paragraph(safe, name_style))
            first_content_line = False
        elif stripped.isupper() and len(stripped) < 60:
            story.append(Paragraph(f'<u>{safe}</u>', heading_style))
        elif stripped.startswith(("•", "-", "*")):
            text = stripped.lstrip("•-* ").strip()
            safe_text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(f"• {safe_text}", bullet_style))
        else:
            story.append(Paragraph(safe, body_style))

    doc.build(story)
    return buf.getvalue()
